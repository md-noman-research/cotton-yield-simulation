"""
build_docx.py
=============
Takes the manuscript markdown file and converts it to a properly formatted
Word document (.docx) for journal submission. Handles headings, inline bold
and italic, tables, and paragraph spacing. Requires python-docx.

Install once with: pip install python-docx
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

MD_PATH  = r"C:\Users\User\Desktop\26_NOBLE_REVISION.md"
OUT_PATH = r"C:\Users\User\Desktop\26_NOBLE_REVISION.docx"

doc = Document()

# ─── Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── Styles
styles = doc.styles

def set_style(style_name, font_name="Times New Roman", size=12, bold=False,
              space_before=0, space_after=6, color=None):
    try:
        style = styles[style_name]
    except KeyError:
        return
    style.font.name   = font_name
    style.font.size   = Pt(size)
    style.font.bold   = bold
    if color:
        style.font.color.rgb = RGBColor(*color)
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after  = Pt(space_after)

set_style("Normal",    size=12)
set_style("Heading 1", size=14, bold=True, space_before=12, space_after=6, color=(0, 0, 0))
set_style("Heading 2", size=13, bold=True, space_before=10, space_after=4)
set_style("Heading 3", size=12, bold=True, space_before=8,  space_after=3)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, italic=False, bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    # Handle inline italic (*text*) and bold (**text**)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if italic:
            run.italic = True
        if bold:
            run.bold = True
    return p

def add_table_from_lines(doc, lines):
    """Parse markdown table lines and add to docx."""
    # filter separator rows
    rows = [l for l in lines if not re.match(r'^\|[-| :]+\|$', l.strip())]
    if not rows:
        return
    
    # parse cells
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)
    
    if not parsed:
        return
    
    n_cols = max(len(r) for r in parsed)
    n_rows = len(parsed)
    
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(parsed):
        row = tbl.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell = row.cells[j]
            # Remove markdown bold/italic
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            cell.text = clean
            # Header row bold
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    
    doc.add_paragraph()  # spacing after table

# ─── Parse markdown and build document
with open(MD_PATH, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
in_table = False
table_lines = []

while i < len(lines):
    raw = lines[i].rstrip('\n')
    stripped = raw.strip()

    # Table detection
    if stripped.startswith('|'):
        if not in_table:
            in_table = True
            table_lines = []
        table_lines.append(stripped)
        i += 1
        continue
    else:
        if in_table:
            add_table_from_lines(doc, table_lines)
            in_table = False
            table_lines = []

    # Headings
    if stripped.startswith('#### '):
        add_heading(doc, stripped[5:], level=4)
    elif stripped.startswith('### '):
        add_heading(doc, stripped[4:], level=3)
    elif stripped.startswith('## '):
        add_heading(doc, stripped[3:], level=2)
    elif stripped.startswith('# '):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stripped[2:])
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)

    # Horizontal rule
    elif stripped in ('---', '***', '___'):
        doc.add_paragraph()

    # Bold line (author, keywords, etc)
    elif stripped.startswith('**') and stripped.endswith('**') and '\n' not in stripped:
        add_paragraph(doc, stripped, bold=True)

    # Empty line
    elif stripped == '':
        pass  # paragraph spacing is handled by space_after

    # Normal paragraph
    else:
        add_paragraph(doc, stripped)

    i += 1

# Flush any remaining table
if in_table and table_lines:
    add_table_from_lines(doc, table_lines)

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
